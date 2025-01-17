import streamlit as st
from azure.core.credentials import AzureKeyCredential
from azure.ai.language.conversations import ConversationAnalysisClient




CW_people = ("Elodie Jaubert","Remi Cazelles","Alban Chesneau","Nicolas CASTET","Lucie Chupin","Mahbub Morshed",
             "Thomas Bottein","Charlotte Gallois","Fabien Bohic","David JUNQUA")
CW_company = dict({"Elodie Jaubert":{'id':1,
                                    'role':'Customer'},
                  "Remi Cazelles":{'id':2,
                                   'role':'Customer'},
                  "Wenbo Li":{'id':3,
                                   'role':'Assistant'},
                  "Alban Chesneau":{'id':4,
                                    'role':'Customer'},
                  "Nicolas CASTET":{'id':5,
                                    'role':'Customer'},
                   "Lucie Chupin":{'id':6,
                                    'role':'Agent'},
                    "Mahbub Morshed":{'id':7,
                                    'role':'Customer'},
                    "Thomas Bottein":{'id':8,
                                    'role':'Customer'},
                    "Charlotte Gallois":{'id':9,
                                    'role':'Agent'},
                    "Fabien Bohic":{'id':10,
                                    'role':'Customer'},
                    
                  })



def assign_role_and_id(line):
    nom = " ".join(line.split(" ")[:2])
    try :
        participant_id = CW_company[nom]["id"]
        role = 'Agent'# CW_company[nom]["role"]
    except :
        participant_id = 99
        role = 'Customer'
       
    return participant_id, role


def format_docx_document(doc):
    formatted_paragraphs=list()
    characters_count=0
    for i, paragraph in enumerate(doc.paragraphs):
        text=list()
        if i <= 2:
            continue
        # Format each paragraph as a dictionary
        for line in paragraph.text.split("\n"):
            
            if line == "":
                continue
            
            elif line.startswith(CW_people):  
                participant_id, role = assign_role_and_id(line)
                

            else:
                if line:
                    text.append(line.replace("\xa0",""))
                    
                    characters_count += len(line.replace("\xa0",""))
        if text:      
                    
            formatted_paragraph = {
                "text": ". ".join(text),
                "id": str(i),  # Assign a unique ID based on the order
                "role": role,      # Assign a role (adjust based on your document content)
                "participantId": participant_id
            }
            
            # Append formatted paragraph to list
            formatted_paragraphs.append(formatted_paragraph)
    return formatted_paragraphs

import unicodedata

def get_text_length(text):
    """Calculate the length of a string in Unicode text elements."""
    return sum(1 for _ in unicodedata.normalize('NFC', text))

def chunk_text_by_length(conversation,max_length=12500):
    """Split the conversation into chunks based on text length in Unicode elements."""
    chunks = []
    current_chunk = []
    current_length = 0

    for item in conversation:
        
        text_length = get_text_length(item['text'])
        if current_length + text_length <= max_length:
            current_chunk.append(item)
            current_length += text_length
        else:
            chunks.append(current_chunk)
            current_chunk = [item]
            current_length = text_length
            
    if current_chunk:
        chunks.append(current_chunk)
               
    return chunks


def view_result(result):
    task_results = result["tasks"]["items"]
    for i,task in enumerate(task_results):
        print(f"\n{task['taskName']} status: {task['status']}")
        task_result = task["results"]
        if task_result["errors"]:
            print("... errors occurred ...")
            for error in task_result["errors"]:
                print(error)
        else:
            conversation_result = task_result["conversations"][0]
            if conversation_result["warnings"]:
                print("... view warnings ...")
                for warning in conversation_result["warnings"]:
                    print(warning)
            else:
                summaries = conversation_result["summaries"]
                for summary in summaries:
                    print(f"{summary['aspect']}: {summary['text']}")
                    st.subheader(f"{summary['aspect']} [{i+1}]:")
                    st.text(f"[{i+1}]: {summary['text']}")





def resume_conv_azure(endpoint="endpoint",key="key",formatted_paragraphs="dict of texts",conversation_id="unique id"):
    
    # chunks = chunk_text_by_length(formatted_paragraphs,max_length=12500)
    # conversation_chunks = chunks

    client = ConversationAnalysisClient(endpoint, AzureKeyCredential(key))
    try:
        results = []
        for idx, chunk in enumerate(formatted_paragraphs):
            
            print(5*"************\n")
            print(idx, len(chunk))
            try:
                conversation_id = f"{conversation_id}_{idx + 1}"
                poller = client.begin_conversation_analysis(
                    task={
                        "displayName": f"Analyze conversations part {idx + 1}",
                        "analysisInput": {
                            "conversations": [
                                {
                                    "conversationItems": chunk,
                                    "modality": "text",
                                    "id": conversation_id,
                                    "language": "fr",
                                },
                            ]
                        },
                        "tasks": [

                            {
                                "taskName": "Title and chapter task",
                                "kind": "ConversationalSummarizationTask",
                                "parameters": {"summaryAspects": ["chapterTitle"]},
                            },
                            {
                                "taskName": "Recap task",
                                "kind": "ConversationalSummarizationTask",
                                "parameters": {"summaryAspects": ["narrative"]},
                            },

                        ],
                    }
                )
                result = poller.result()
                results.append(result)
                view_result(result)
            except Exception as e:
                print(f"Error in chunk {idx + 1}: {e}")
            continue
            
            
                
                
    except Exception as e:
        print(e)
        
    finally:
        print("polling finishes")
        client.close()
        
    return results
                        
                        
                        
                        
from io import BytesIO
from docx import Document

def document_to_bytes(doc: Document) -> BytesIO:
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  
    return buffer


def create_document(results,nom_discussion,uuid):
    recap = Document()

    recap.add_heading(f'Résumé : {nom_discussion}')
    recap.add_paragraph("Ce resume utilise l'outil Microsoft Azure ConversationAnalysisClient avec les paramètres 'summaryAspects': [chapterTitle] & [narrative]")
    recap.add_paragraph(f"UUID de résumé :{uuid}")
    
    for i, result in enumerate(results):
        task_results = result["tasks"]["items"]
        for j, task in enumerate(task_results):
            recap.add_heading(f"\n{task['taskName']} status {j+1}: {task['status']}", level=2)
            task_result = task["results"]
            if task_result["errors"]:
                print("... errors occurred ...")
                for error in task_result["errors"]:
                    print(error)
            else:
                conversation_result = task_result["conversations"][0]
                if conversation_result["warnings"]:
                    print("... view warnings ...")
                    for warning in conversation_result["warnings"]:
                        print(warning)
                else:
                    summaries = conversation_result["summaries"]
                    for summary in summaries:
                        recap.add_paragraph(f"{summary['aspect']} {j+1}: {summary['text']}")
                        
                        
    return document_to_bytes(recap)
    
