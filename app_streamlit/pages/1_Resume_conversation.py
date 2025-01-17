import streamlit as st
import docx
from utils.helper_functions import format_docx_document, resume_conv_azure, chunk_text_by_length, view_result, create_document
import uuid
from azure.identity import DefaultAzureCredential, ManagedIdentityCredential
from azure.keyvault.secrets import SecretClient
from dotenv import load_dotenv
import os
import hmac

load_dotenv()

# Azure Key Vault details
KEY_VAULT_NAME = os.getenv('KEY_VAULT_NAME')
SECRET_NAME = os.getenv('SECRET_NAME')
KV_URI = f"https://{KEY_VAULT_NAME}.vault.azure.net"


st.set_page_config(page_title="Resume_conversation", page_icon=":speech_balloon:")


# Authenticate and retrieve the password from Azure Key Vault
@st.cache_data
def get_password_from_keyvault():
    """Fetches the password securely from Azure Key Vault."""
    credential = DefaultAzureCredential() # for development ONLY
    # credential = ManagedIdentityCredential()
    client = SecretClient(vault_url=KV_URI, credential=credential)
    retrieved_secret = client.get_secret(SECRET_NAME)
    return retrieved_secret.value

# Fetch the password once
STORED_PASSWORD = get_password_from_keyvault()

# STORED_PASSWORD = os.getenv('STORED_PASSWORD')
def check_password():
    """Returns `True` if the user had the correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if hmac.compare_digest(st.session_state["password"], STORED_PASSWORD):
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store the password.
        else:
            st.session_state["password_correct"] = False

    # Return True if the password is validated.
    if st.session_state.get("password_correct", False):
        return True

    # Show input for password.
    st.text_input(
        "Password", type="password", on_change=password_entered, key="password"
    )
    if "password_correct" in st.session_state:
        st.error("😕 Password incorrect")
    return False



if not check_password():
    st.stop()  # Do not continue if check_password is not True.

st.success("🎉 You have access to the app!")



LANGUAGE_KEY = os.getenv('LANGUAGE_KEY')
LANGUAGE_ENDPOINT = os.getenv('LANGUAGE_ENDPOINT')



# app starts from here
st.markdown("# Résumé conversation")
st.sidebar.header("Résumé conversation")

st.write(
    """Ici vous pouvez charge un ficher transcript d'un meeting TEAMS, la conversation sera archivée ainsi 
    que son résumé. Résumé généré automatiquement par OpenAI à traver l'outil Azure azure.ai.language.conversations.
    La limite de taille du fichier est 200 MB"""
)


uploaded_file  = st.file_uploader(label='Fichier de conversation à analyser',
                 type=['docx'],
                 accept_multiple_files=False,
                 help='Vous trouverez les transcripts des réunions TEAMS sur : [Y:\10- Administratif\2- Comptes-rendus\2024]')

# progress_bar = st.sidebar.progress(0)
# status_text = st.sidebar.empty()
# last_rows = np.random.randn(1, 1)
# chart = st.line_chart(last_rows)


if uploaded_file is not None:
    doc = docx.Document(uploaded_file)
    formatted_paragraphs = format_docx_document(doc)
    st.success('Document chargé avec succés')
    
    chunks = chunk_text_by_length(formatted_paragraphs,max_length=25000)
    st.text(f"Le texte d'origine est segmenté en {len(chunks)} parties contenant {[len(a) for a in chunks]} interventions chacune")

    nom_discussion = st.text_input('Définir un nom pour le résumé de conversation',
                  placeholder="ex. 2024_12_25_Discussion_Petit_Pilote")
    
    st.caption("Remplissez le champ ci-dessus et appuyez sur entrée pour continuer")
    if nom_discussion:
        
        unique_conversation_id = str(uuid.uuid1())
        st.text(f"Le fichier {nom_discussion}.docx sera disponible pour téléchargement après avoir cliqué sur 'Obtenir résumé")
        st.caption(" Le résumé est sauvegardé avec un UUID (Universal Unique Identifier) de 128 bits. Il est unique car généré sur le temps actuel est la machine utilisée")
        

        if st.button('Obtenir résumé'):
            
            results = resume_conv_azure(endpoint=LANGUAGE_ENDPOINT,
                                       key=LANGUAGE_KEY,
                                       formatted_paragraphs=chunks,
                                       conversation_id=unique_conversation_id)
            
            docx_buffer = create_document(results,nom_discussion,unique_conversation_id)
            
            st.download_button(
                label='Download Resume',
                data=docx_buffer,
                file_name=f"{nom_discussion}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # for res in results:
            #     view_result(res)
                
                

            
            # export_doc = docx.Document()
            # export_doc.add_heading(f'{nom_discussion}',0)
            # export_doc.add_heading(f'Résumé automatique',1)
            # export_doc.add_paragraph(f'Conversation UUID : {unique_conversation_id}')
            # export_doc.add_heading(f'Participants : {participants}', level=2)          
  
                            
            # # CHANGE PATH HERE FOR DEVELOPMENT
            # # export_doc.save(f'../Resumes_doc/{nom_discussion}.docx')
            # export_doc.save(f'/app_streamlit/Resume_doc/{nom_discussion}.docx')
                            
                            
            
    
